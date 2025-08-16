import os
from datetime import datetime
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import uuid

class EnhancedDocumentGenerator:
    def __init__(self):
        self.universities = {
            "harvard": {
                "name": "Harvard University",
                "address": "Cambridge, MA 02138",
                "phone": "(617) 495-1000",
                "website": "harvard.edu",
                "color": "8B0000"  # Dark red
            },
            "stanford": {
                "name": "Stanford University", 
                "address": "Stanford, CA 94305",
                "phone": "(650) 723-2300",
                "website": "stanford.edu",
                "color": "CC0000"  # Cardinal red
            },
            "mit": {
                "name": "Massachusetts Institute of Technology",
                "address": "Cambridge, MA 02139",
                "phone": "(617) 253-1000",
                "website": "mit.edu",
                "color": "8A2BE2"  # Blue violet
            },
            "generic": {
                "name": "University",
                "address": "University Address",
                "phone": "University Phone",
                "website": "university.edu",
                "color": "2F4F4F"  # Dark slate gray
            }
        }
        
        self.document_templates = {
            "bonafide": {
                "title": "BONAFIDE CERTIFICATE",
                "template": "This is to certify that {student_name}, Roll No: {roll_number}, is a bonafide student of this institution. He/She has been studying in this institution since {admission_date} in the {course} program in the Department of {department}. This certificate is issued for the purpose of {purpose}."
            },
            "noc": {
                "title": "NO OBJECTION CERTIFICATE",
                "template": "This is to certify that we have no objection to {student_name}, Roll No: {roll_number}, a student of {course} program in the Department of {department}, for {purpose}. The student is in good academic standing and has no disciplinary issues."
            },
            "character": {
                "title": "CHARACTER CERTIFICATE",
                "template": "This is to certify that {student_name}, Roll No: {roll_number}, is a student of {course} program in the Department of {department}. During his/her stay in this institution, his/her character and conduct have been found to be satisfactory. This certificate is issued for {purpose}."
            },
            "transfer": {
                "title": "TRANSFER CERTIFICATE",
                "template": "This is to certify that {student_name}, Roll No: {roll_number}, was a student of this institution from {admission_date} studying {course} in the Department of {department}. He/She is now seeking transfer for {purpose}. His/Her character and conduct during the stay were satisfactory."
            },
            "fee_structure": {
                "title": "FEE STRUCTURE LETTER",
                "template": "This letter provides the official fee structure for {student_name}, Roll No: {roll_number}, enrolled in {course} program in the Department of {department}. This information is provided for {purpose}. Please contact the finance office for detailed fee breakdown."
            },
            "transcript": {
                "title": "ACADEMIC TRANSCRIPT REQUEST",
                "template": "This acknowledges the request for official academic transcripts for {student_name}, Roll No: {roll_number}, enrolled in {course} program in the Department of {department}. The transcripts are requested for {purpose}. Official transcripts will be processed within 5-7 business days."
            }
        }

    def generate_unique_filename(self, document_type, student_name, purpose):
        """Generate a unique filename with timestamp and UUID"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        clean_name = student_name.replace(" ", "_").replace("/", "_")
        clean_purpose = purpose.replace(" ", "_").replace("/", "_")[:20]
        
        return f"{document_type}_{clean_name}_{clean_purpose}_{timestamp}_{unique_id}.docx"

    def add_watermark(self, doc, text):
        """Add a watermark to the document"""
        section = doc.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Make watermark semi-transparent
        run = paragraph.runs[0]
        run.font.size = Pt(60)
        run.font.color.rgb = None  # This will be set to light gray
        run.font.name = 'Arial'

    def add_university_logo_placeholder(self, doc, university_name):
        """Add a placeholder for university logo"""
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = paragraph.add_run(f"[{university_name} LOGO]")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = None

    def create_document_header(self, doc, university_data):
        """Create the document header with university information"""
        # University name
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(university_data["name"].upper())
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.name = 'Times New Roman'
        
        # Office subtitle
        office_para = doc.add_paragraph()
        office_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        office_run = office_para.add_run("Office of the Registrar")
        office_run.font.size = Pt(16)
        office_run.font.name = 'Times New Roman'
        
        # Address and contact info
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(f"{university_data['address']}\nTel: {university_data['phone']}")
        contact_run.font.size = Pt(12)
        contact_run.font.name = 'Times New Roman'
        
        # Add a line separator
        doc.add_paragraph().add_run("_" * 80)

    def create_document_title(self, doc, title):
        """Create the document title"""
        doc.add_paragraph()  # Add some space
        
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_run.font.name = 'Times New Roman'
        
        # Add underline
        title_run.font.underline = True
        
        doc.add_paragraph()  # Add some space

    def create_document_body(self, doc, template_text, student_data):
        """Create the main document body"""
        body_para = doc.add_paragraph()
        body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_run = body_para.add_run(template_text)
        body_run.font.size = Pt(12)
        body_run.font.name = 'Times New Roman'
        body_run.font.space_after = Pt(12)

    def create_document_footer(self, doc, university_data, student_data):
        """Create the document footer with date, place, and signature"""
        doc.add_paragraph()  # Add space
        
        # Date and place
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        current_date = datetime.now().strftime("%B %d, %Y")
        date_run = date_para.add_run(f"Date: {current_date}")
        date_run.font.size = Pt(12)
        date_run.font.name = 'Times New Roman'
        
        place_para = doc.add_paragraph()
        place_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        place_run = place_para.add_run(f"Place: {university_data['name']}")
        place_run.font.size = Pt(12)
        place_run.font.name = 'Times New Roman'
        
        # Signature section
        doc.add_paragraph()  # Add space
        
        signature_para = doc.add_paragraph()
        signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Signature line
        signature_run = signature_para.add_run("_" * 40)
        signature_run.font.size = Pt(12)
        signature_run.font.name = 'Times New Roman'
        
        # Registrar title
        registrar_para = doc.add_paragraph()
        registrar_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        registrar_run = registrar_para.add_run("Registrar")
        registrar_run.font.size = Pt(12)
        registrar_run.font.bold = True
        registrar_run.font.name = 'Times New Roman'
        
        # University seal
        seal_para = doc.add_paragraph()
        seal_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        seal_run = seal_para.add_run(university_data['name'])
        seal_run.font.size = Pt(10)
        seal_run.font.name = 'Times New Roman'

    def add_document_metadata(self, doc, student_data, document_type):
        """Add metadata to the document"""
        core_props = doc.core_properties
        core_props.title = f"{document_type.title()} - {student_data['student_name']}"
        core_props.subject = f"University Document: {document_type}"
        core_props.author = "University Registrar Office"
        core_props.comments = f"Generated for {student_data['student_name']} on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"

    def generate_document(self, university_code, document_type, student_data):
        """Generate a complete document with all components"""
        # Validate inputs
        if university_code not in self.universities:
            raise ValueError(f"Invalid university code: {university_code}")
        
        if document_type not in self.document_templates:
            raise ValueError(f"Invalid document type: {document_type}")
        
        # Get university and template data
        university_data = self.universities[university_code]
        template_data = self.document_templates[document_type]
        
        # Create new document
        doc = Document()
        
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add watermark
        self.add_watermark(doc, "OFFICIAL DOCUMENT")
        
        # Add university logo placeholder
        self.add_university_logo_placeholder(doc, university_data["name"])
        
        # Create document header
        self.create_document_header(doc, university_data)
        
        # Create document title
        self.create_document_title(doc, template_data["title"])
        
        # Format student name with title
        formatted_name = f"Mr./Ms. {student_data['student_name']}"
        
        # Create document body
        body_text = template_data["template"].format(
            student_name=formatted_name,
            roll_number=student_data['roll_number'],
            course=student_data['course'],
            department=student_data['department'],
            admission_date=student_data['admission_date'],
            purpose=student_data['purpose']
        )
        self.create_document_body(doc, body_text, student_data)
        
        # Create document footer
        self.create_document_footer(doc, university_data, student_data)
        
        # Add document metadata
        self.add_document_metadata(doc, student_data, document_type)
        
        # Generate unique filename
        filename = self.generate_unique_filename(document_type, student_data['student_name'], student_data['purpose'])
        
        return doc, filename

    def save_document(self, doc, filename, output_dir="generated"):
        """Save the document to the specified directory"""
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        # Full path for the file
        file_path = os.path.join(output_dir, filename)
        
        # Save the document
        doc.save(file_path)
        
        return file_path

    def generate_and_save(self, university_code, document_type, student_data, output_dir="generated"):
        """Generate and save a document in one step"""
        doc, filename = self.generate_document(university_code, document_type, student_data)
        file_path = self.save_document(doc, filename, output_dir)
        
        return {
            "success": True,
            "filename": filename,
            "file_path": file_path,
            "download_url": f"/download/{filename}",
            "message": f"Document generated successfully: {filename}"
        }

# Example usage and testing
if __name__ == "__main__":
    generator = EnhancedDocumentGenerator()
    
    # Sample student data
    sample_student = {
        "student_name": "John Smith",
        "roll_number": "CS2024001",
        "course": "Bachelor of Science (B.Sc.)",
        "department": "Computer Science",
        "admission_date": "2023-08-15",
        "purpose": "Internship application"
    }
    
    # Generate a sample document
    try:
        result = generator.generate_and_save("harvard", "bonafide", sample_student)
        print(f"Document generated: {result['filename']}")
        print(f"File saved at: {result['file_path']}")
    except Exception as e:
        print(f"Error generating document: {e}")
